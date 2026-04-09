import hashlib
import logging
import json
import os
from datetime import datetime, timezone

from docx import Document
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger(__name__)

OUTPUT_DIR = os.getenv("OUTPUT_DIR", "./output")


def _generate_chunk_id(file_path: str, text: str) -> str:
    """Generate a unique chunk ID using SHA256 hash."""
    return hashlib.sha256(
        (file_path + text).encode("utf-8")
    ).hexdigest()[:16]


def _save_blocks(blocks: list[dict]) -> None:
    """Append blocks to scraped.jsonl output file."""
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, "scraped.jsonl")
    try:
        with open(output_path, "a", encoding="utf-8") as f:
            for block in blocks:
                f.write(
                    json.dumps(block, ensure_ascii=False) + "\n"
                )
        logger.info(f"Saved {len(blocks)} blocks to {output_path}")
    except Exception as e:
        logger.error(f"Failed to save blocks: {e}")


def _extract_table_text(table) -> str:
    """Flatten a Word table into readable text."""
    rows_text = []

    # Get header row
    headers = []
    if table.rows:
        headers = [
            cell.text.strip()
            for cell in table.rows[0].cells
        ]

    # Process data rows
    for row_num, row in enumerate(table.rows):
        if row_num == 0:
            continue
        row_parts = []
        for header, cell in zip(headers, row.cells):
            value = cell.text.strip()
            if value:
                if header:
                    row_parts.append(f"{header}: {value}")
                else:
                    row_parts.append(value)
        if row_parts:
            rows_text.append(
                f"Row {row_num}: " + ", ".join(row_parts)
            )

    return "\n".join(rows_text)


def _get_heading_level(paragraph) -> str:
    """Get heading level from paragraph style name."""
    style_name = paragraph.style.name.lower()
    if "heading 1" in style_name:
        return "h1"
    elif "heading 2" in style_name:
        return "h2"
    elif "heading 3" in style_name:
        return "h3"
    elif "heading 4" in style_name:
        return "h4"
    return ""


def load_word(
    file_path    : str,
    document_type: str,
    org_name     : str
) -> list[dict]:
    """
    Load and extract text from a Word (.docx) file.
    Preserves heading structure and extracts tables.
    Returns list of block dicts.
    """
    logger.info(f"Loading Word document: {file_path}")

    if not os.path.exists(file_path):
        logger.error(f"Word file not found: {file_path}")
        return []

    document_title = os.path.basename(file_path)
    extracted_at   = datetime.now(timezone.utc).isoformat()
    blocks         = []

    try:
        doc = Document(file_path)
    except Exception as e:
        logger.error(
            f"Failed to open Word file {file_path}: {e}"
        )
        return []

    # Section context trackers
    section_title = ""
    section_level = ""
    chapter       = ""
    article       = ""

    # Walk document elements in order
    # python-docx iterates paragraphs and tables separately
    # We use the document body XML to preserve order
    from docx.oxml.ns import qn

    body = doc.element.body
    
    # Pre-build hash maps for elements O(1) lookup
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body.iterchildren():
        tag = child.tag

        # Paragraph
        if tag == qn("w:p"):
            # O(1) hashmap lookup
            para = para_map.get(child)
            if not para:
                continue
                
            para_text  = para.text.strip()
            style_name = para.style.name.lower()

            if not para_text:
                continue

            # Determine if heading
            is_heading = "heading" in style_name
            heading_level = ""

            if "heading 1" in style_name:
                heading_level = "h1"
            elif "heading 2" in style_name:
                heading_level = "h2"
            elif "heading 3" in style_name:
                heading_level = "h3"
            elif "heading 4" in style_name:
                heading_level = "h4"

            if is_heading and heading_level:
                content_type  = "heading"
                section_title = para_text
                section_level = heading_level

                if heading_level in ["h1", "h2"]:
                    chapter = para_text
                    article = ""
                else:
                    article = para_text

            else:
                content_type = "paragraph"

            chunk_id = _generate_chunk_id(file_path, para_text)

            blocks.append({
                "chunk_id"      : chunk_id,
                "source_url"    : file_path,
                "document_title": document_title,
                "organization"  : org_name,
                "section_title" : section_title,
                "section_level" : section_level,
                "chapter"       : chapter,
                "article"       : article,
                "content_type"  : content_type,
                "text"          : para_text,
                "char_count"    : len(para_text),
                "extracted_at"  : extracted_at,
            })

        # Table
        elif tag == qn("w:tbl"):
            table = table_map.get(child)
            if table:
                try:
                    table_text = _extract_table_text(table)
                    if not table_text.strip():
                        continue

                    chunk_id = _generate_chunk_id(
                        file_path, table_text
                    )
                    blocks.append({
                        "chunk_id"      : chunk_id,
                        "source_url"    : file_path,
                        "document_title": document_title,
                        "organization"  : org_name,
                        "section_title" : section_title,
                        "section_level" : section_level,
                        "chapter"       : chapter,
                        "article"       : article,
                        "content_type"  : "table",
                        "text"          : table_text,
                        "char_count"    : len(table_text),
                        "extracted_at"  : extracted_at,
                    })
                except Exception as e:
                    logger.warning(f"Table extraction failed: {e}")
                    continue

    _save_blocks(blocks)

    logger.info(
        f"Word loading complete. "
        f"Blocks extracted: {len(blocks)} from {file_path}"
    )
    return blocks